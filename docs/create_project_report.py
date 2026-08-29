from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Project_FORESIGHT_Report.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 110, 128)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        shade(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, 9.5, NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_font(r, 9.5)
    table.style = "Table Grid"
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("PROJECT FORESIGHT  |  Technical Project Report")
    set_font(r, 9, MUTED, bold=True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Project FORESIGHT - Demand & Inventory Intelligence Platform")
    set_font(r, 8.5, MUTED)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT REPORT")
    set_font(r, 11, BLUE, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT FORESIGHT")
    set_font(r, 30, NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Demand Forecasting and Inventory Intelligence")
    set_font(r, 15, RGBColor(43, 81, 99))
    p.paragraph_format.space_after = Pt(34)
    table = add_table(doc, ["Project", "Technology", "Links"], [[
        "Retail supply-chain decision support platform",
        "Next.js, FastAPI, SQLAlchemy, scikit-learn, PostgreSQL / SQLite",
        "GitHub: github.com/shaikmuzamil1844-coder/project_foresight\nLive app: project-foresight-frontend.vercel.app",
    ]], [3000, 3000, 3360])
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Prepared for project evaluation | {date.today().strftime('%d %B %Y')}")
    set_font(r, 10, MUTED, italic=True)

    doc.add_page_break()
    heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Project FORESIGHT is a full-stack retail intelligence platform that converts historical sales data into actionable inventory decisions. "
        "It enables analysts to import transaction data, monitor SKU-level risk, generate demand forecasts, and prioritize replenishment activity from one dashboard."
    )
    heading(doc, "Business Problem", 2)
    bullet(doc, "Stockouts cause lost sales and weaken customer trust when demand exceeds available inventory.")
    bullet(doc, "Overstock locks cash in slow-moving goods and increases storage and obsolescence costs.")
    bullet(doc, "Spreadsheet-based replenishment cannot consistently account for demand variability, history, and supplier lead time.")
    heading(doc, "Solution", 2)
    doc.add_paragraph(
        "FORESIGHT combines data ingestion, machine-learning forecasts, statistical safety-stock calculations, and an interactive web interface. "
        "After a dataset is uploaded or seeded, the platform stores the records, recalculates recommendations, and serves the resulting KPIs, charts, and forecasts through a FastAPI API."
    )

    heading(doc, "2. System Architecture")
    add_table(doc, ["Layer", "Responsibility", "Implementation"], [
        ("Frontend", "Dashboard, inventory matrix, upload flow, forecasts, and assistant UI", "Next.js 16, React 19, TypeScript, Recharts"),
        ("API", "Typed REST endpoints and request validation", "FastAPI and Pydantic"),
        ("Data", "Relational storage of products, sales, inventory, forecasts, and recommendations", "SQLAlchemy with SQLite locally or PostgreSQL in production"),
        ("Analytics", "Feature engineering, demand forecasts, and replenishment calculation", "pandas, NumPy, scikit-learn GradientBoostingRegressor"),
        ("AI assistant", "Natural-language inventory guidance grounded in stored recommendations", "Gemini when configured, with deterministic local responses otherwise"),
    ], [1500, 3900, 3960])
    heading(doc, "Request Flow", 2)
    doc.add_paragraph("User action -> Next.js client -> FastAPI router -> database and analytics services -> JSON response -> dashboard visualization.")

    heading(doc, "3. Data Workflow and API")
    doc.add_paragraph(
        "The CSV/XLSX import accepts retail transaction fields including date, SKU identifier, product name, category, units sold, and price. "
        "During ingestion, products are upserted, transaction and inventory records are refreshed, and each SKU receives a risk calculation and available forecast."
    )
    add_table(doc, ["Endpoint", "Purpose"], [
        ("POST /api/upload/csv", "Validate and ingest a user dataset."),
        ("POST /api/upload/seed", "Load the included retail sample dataset."),
        ("GET /api/dashboard/summary", "Return data-derived KPIs and risk counts."),
        ("GET /api/inventory/risk-matrix", "Return current SKU-level inventory risk."),
        ("GET /api/forecast/{sku_id}?days=7-90", "Generate a bounded horizon forecast."),
        ("POST /api/assistant/query", "Return contextual inventory guidance."),
    ], [4000, 5360])

    heading(doc, "4. Forecasting and Inventory Logic")
    heading(doc, "Demand Forecasting", 2)
    doc.add_paragraph(
        "The forecasting service derives calendar features, lagged demand values (1, 7, 14, and 28 days), and rolling means and standard deviations. "
        "A Gradient Boosting regressor is evaluated on a time-based holdout set and then recursively forecasts the requested horizon. "
        "The API returns MAE, RMSE, MAPE, and 95% confidence bounds based on model error."
    )
    heading(doc, "Inventory Formulas", 2)
    add_table(doc, ["Measure", "Formula / Rule"], [
        ("Lead-time demand", "Average daily demand x supplier lead time"),
        ("Safety stock", "max(minimum safety stock, 1.65 x demand standard deviation x sqrt(lead time))"),
        ("Reorder point", "Lead-time demand + safety stock"),
        ("Replenishment quantity", "max(0, reorder point + 30 days of demand - current stock)"),
        ("Risk level", "High when stock is at/below ROP or stockout is within lead time; medium near ROP; overstock above 60 days of demand"),
    ], [2800, 6560])

    heading(doc, "5. Demonstrated Result")
    doc.add_paragraph(
        "An end-to-end verification was performed using the included sample data. The seed operation stored 3,660 records across 10 SKUs. "
        "The dashboard then returned data-derived totals, the inventory API returned recalculated risk records, and a 7-day SKU forecast returned 14 historical observations plus 7 forecast days."
    )
    heading(doc, "6. Deployment and Security")
    bullet(doc, "Frontend: Vercel deployment at https://project-foresight-frontend.vercel.app")
    bullet(doc, "Backend: Render deployment configured through render.yaml.")
    bullet(doc, "Configuration is supplied through DATABASE_URL, GEMINI_API_KEY, and CORS_ORIGINS environment variables; secrets are not embedded in application source.")
    bullet(doc, "CORS is restricted to configured origins, and forecast horizons are validated between 7 and 90 days.")

    heading(doc, "7. Conclusion")
    doc.add_paragraph(
        "Project FORESIGHT demonstrates how operational retail data can become a practical replenishment workflow. "
        "Its modular architecture separates the presentation layer, API, persistence, forecasting, and inventory logic, making it suitable for future enhancements such as supplier-specific lead times, authentication, scheduled retraining, and richer scenario planning."
    )

    doc.core_properties.title = "Project FORESIGHT - Technical Project Report"
    doc.core_properties.author = "Project FORESIGHT Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
